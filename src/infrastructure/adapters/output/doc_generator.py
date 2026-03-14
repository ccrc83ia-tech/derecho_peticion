from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from src.domain.models import Branding
from src.domain.ports.out_ports import FileExporterPort

OUTPUT_DIR = Path("generated_docs")


class DocxEngine(FileExporterPort):

    def __init__(self, output_dir: Path = OUTPUT_DIR) -> None:
        self._output_dir = output_dir
        self._output_dir.mkdir(parents=True, exist_ok=True)

    async def export(self, content: str, branding: Branding, file_name: str) -> str:
        doc = Document()

        # Header with branding
        if branding.header_text:
            header_para = doc.add_paragraph()
            run = header_para.add_run(branding.header_text)
            run.bold = True
            run.font.size = Pt(14)
            try:
                r, g, b = self._hex_to_rgb(branding.primary_color)
                run.font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                pass

        doc.add_paragraph("")  # spacer

        # Body content — each paragraph from AI
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        # Footer
        if branding.footer_text:
            doc.add_paragraph("")
            footer_run = doc.add_paragraph().add_run(branding.footer_text)
            footer_run.font.size = Pt(8)
            footer_run.italic = True

        file_path = self._output_dir / file_name
        doc.save(str(file_path))
        return file_name

    @staticmethod
    def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
        h = hex_color.lstrip("#")
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
